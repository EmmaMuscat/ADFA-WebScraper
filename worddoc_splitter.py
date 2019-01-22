import docx

def print_paragraphs(filename):
    doc = docx.Document(filename)
    for j in range(len(doc.paragraphs)):
        print(doc.paragraphs[j].text)
        print(len(doc.paragraphs[j].runs))
        print('\n')

def show_breaks(filename):
    doc = docx.Document(filename)
    dict = {}
    #check each paragraph
    for i in range(len(doc.paragraphs)):

        if len(doc.paragraphs[i].runs) == 0:
            print("skipped! \n")
            continue

        for j in range(len(doc.paragraphs[i].runs)):
            print(doc.paragraphs[j].runs[j].text)

            if(len(doc.paragraphs[i].runs[j]._element.br_lst)==0):
                print("skip 2! \n")
                continue

            for k in range(len(doc.paragraphs[i].runs[j]._element.br_lst)):
                    print(doc.paragraphs[i].runs[j].br_lst[k].type)

def seperate_attempt_3(filename):
    doc = docx.Document(filename)
    dict = {}
    heading = ["" for x in range(len(doc.paragraphs))]
    head_count = 0
    #for each paragraph
    for i in range(len(doc.paragraphs)):
        #check if current paragraph is empty
        if len(doc.paragraphs[i].runs) == 0:
            continue

        #its a heading
        elif (doc.paragraphs[i].runs[0].bold):
            head_count = head_count + 1
            for j in range(len(doc.paragraphs[i].runs)):
                if(doc.paragraphs[i].runs[j].bold):
                    heading[head_count] = heading[head_count] + doc.paragraphs[i].runs[j].text
                else: #if the rest of it isnt bold its text!
                    key = heading[head_count]
                    if key in dict:
                        dict[key] = dict.get(key) + doc.paragraphs[i].text
                    else:
                        dict[key] = doc.paragraphs[i].text

            print(heading[head_count])
            print("\n")
        else:
            #no its just text that isnt empty add to its heading
            key = heading[head_count]
            if key in dict:
                dict[key] = dict.get(key) + doc.paragraphs[i].text
            else:
                dict[key] = doc.paragraphs[i].text
    for a in dict:
        print("KEY")
        print(a)
        print('\n')
        print('\n')
        print("TEXT")
        print(dict[a])
        print('\n')
        print('\n')

def seperate_works_2(filename):
    doc = docx.Document(filename)
    dict = {}
    heading = ["" for x in range(len(doc.paragraphs))]
    head_count = -1
    main_heading = ["" for x in range(len(doc.paragraphs))]
    main_h_count = -1

    #for each paragraph
    for i in range(len(doc.paragraphs)):

        #check if current paragraph is empty
        if len(doc.paragraphs[i].runs) == 0:
            continue

        #its a heading
        elif (doc.paragraphs[i].runs[0].bold):
            if (doc.paragraphs[i].runs[0].text.startswith('V')):
                #it is a big heading
                main_h_count = main_h_count + 1
                for j in range(len(doc.paragraphs[i].runs)):
                    if(doc.paragraphs[i].runs[j].bold):
                        main_heading[main_h_count] = main_heading[main_h_count] + doc.paragraphs[i].runs[j].text
                main_heading[main_h_count] = main_heading[main_h_count] + ":"



            else:
                #its a smol heading
                head_count = head_count + 1
                for j in range(len(doc.paragraphs[i].runs)):
                    if(doc.paragraphs[i].runs[j].bold):
                        heading[head_count] = heading[head_count] + doc.paragraphs[i].runs[j].text
                    else: #if the rest of it isnt bold its text!
                        key = main_heading[main_h_count] + heading[head_count]
                        if key in dict:
                            dict[key] = dict.get(key) + doc.paragraphs[i].text
                        else:
                            dict[key] = doc.paragraphs[i].text



        else:
            #no its just text that isnt empty add to its  with a main heading for uniqueness
            key =  main_heading[main_h_count] + heading[head_count]
            if key in dict:
                dict[key] = dict.get(key) + doc.paragraphs[i].text
            else:
                dict[key] = doc.paragraphs[i].text

    for a in dict:
        print("KEY")
        print(a)
        print('\n')
        print('\n')
        print("TEXT")
        print(dict[a])
        print('\n')
        print('\n')

def seperate_works_1(filename):
    doc = docx.Document(filename)
    dict = {}
    heading = ["" for x in range(len(doc.paragraphs))]
    head_count = 0
    main_heading = ["" for x in range(len(doc.paragraphs))]
    main_h_count = 0

    #for each paragraph
    for i in range(len(doc.paragraphs)):
        #check if current paragraph is empty
        if len(doc.paragraphs[i].runs) == 0:
            continue

        #its a heading
        elif (doc.paragraphs[i].runs[0].bold):
            #it is a big heading
            #print(doc.paragraphs[i].runs[0].font.size)

            #is currently document specific re font size
            #should work out how to determine this dynamically
            if (doc.paragraphs[i].runs[0].font.size == 152400):
                main_h_count = main_h_count + 1
                for j in range(len(doc.paragraphs[i].runs)):
                    if(doc.paragraphs[i].runs[j].bold):
                        main_heading[main_h_count] = main_heading[main_h_count] + doc.paragraphs[i].runs[j].text
                main_heading[main_h_count] = main_heading[main_h_count] + ":"


            #its a smol heading
            else:
                head_count = head_count + 1
                for j in range(len(doc.paragraphs[i].runs)):
                    if(doc.paragraphs[i].runs[j].bold):
                        heading[head_count] = heading[head_count] + doc.paragraphs[i].runs[j].text
                    else: #if the rest of it isnt bold its text!
                        key = main_heading[main_h_count] + heading[head_count]
                        if key in dict:
                            dict[key] = dict.get(key) + doc.paragraphs[i].text
                        else:
                            dict[key] = doc.paragraphs[i].text



        else:
            #no its just text that isnt empty add to its  with a main heading for uniqueness
            key =  main_heading[main_h_count] + heading[head_count]
            if key in dict:
                dict[key] = dict.get(key) + doc.paragraphs[i].text
            else:
                dict[key] = doc.paragraphs[i].text

    for a in dict:
        print("KEY")
        print(a)
        print('\n')
        print('\n')
        print("TEXT")
        print(dict[a])
        print('\n')
        print('\n')

def seperate_attempt_2(filename):
    doc = docx.Document(filename)
    dict = {}
    low_head = []
    high_head = []

    #for each paragraph
    for i in range(len(doc.paragraphs)):
        #check if current paragraph is empty
        if len(doc.paragraphs[i].runs) == 0:
            continue

        #its a heading
        if (doc.paragraphs[i].runs[0].bold):
            #we have a new heading


            if (len(doc.paragraphs[i+1].runs) != 0 ):
                for j in range(len(doc.paragraphs[i+1].runs)):
                    #check if high or low level head and add accordingly
                    if (doc.paragraphs[i+1].runs[j].bold):
                        high_head.append(' ')
                        high_head.append(doc.paragraphs[i].runs[j].text)


            else:
                #check if its bold for each run
                for j in range(len(doc.paragraphs[i].runs)):
                    #check if high or low level head and add accordingly
                    if (doc.paragraphs[i].runs[j].bold):
                        low_head.append(' ')
                        low_head.append(doc.paragraphs[i].runs[j].text)
                print(low_head[i])

        else:
            #no its just text that isnt empty add to its heading
            key = high_head + low_head
            if key in dict:
                dict[key] = dict.get(key) + doc.paragraphs[i].text
            else:
                dict[key] = doc.paragraphs[i].text




    for a in dict:
        print('\n')
        print('\n')
        print('\n')
        print('\n')
        print("KEY")
        print(a)
        print('\n')
        print('\n')
        print("TEXT")
        print(dict[a])
        print('\n')
        print('\n')

def seperate_attempt_1(filename):
    doc = docx.Document(filename)
    dict = {}

    #for each paragraph
    for i in range(len(doc.paragraphs)):
        #check if current paragraph is empty
        if len(doc.paragraphs[i].runs) == 0:
            continue
        #check if current paragraph is a bolded heading
        elif doc.paragraphs[i].runs[0].bold:
            #yes it is check next paragraph for boldness and runs
            if(len(doc.paragraphs[i+1].runs) != 0):
                if doc.paragraphs[i+1].runs[0].bold:
                    #we got a  high level heading
                    high_head = doc.paragraphs[i].runs[0].text
                else:
                    #its a low level heading next one aint bold
                    low_head = doc.paragraphs[i].runs[0].text

        else:
            #no its just text that isnt empty check if key exists
            key = high_head + '_' + low_head
            print("KEY is: ", key)
            if key in dict:
                dict[key] = dict.get(key) + doc.paragraphs[i].text
            else:
                dict[key] = doc.paragraphs[i].text


    for a in dict:
        print("KEY")
        print(a)
        print('\n')
        print('\n')
        print("TEXT")
        print(dict[a])
        print('\n')
        print('\n')

def show_runs(filename):
    doc = docx.Document(filename)
    dict = {}
    #check each paragraph
    for i in range(len(doc.paragraphs)):
        print(len(doc.paragraphs[i].runs))
        print(doc.paragraphs[i].text)

def show_font_size(filename):
    doc = docx.Document(filename)
    dict = {}
    #check each paragraph
    for i in range(len(doc.paragraphs)):
        for j in range(len(doc.paragraphs[i].runs)):
            if(doc.paragraphs[i].runs[j].font.size is not None):
                print(doc.paragraphs[i].runs[j].font.size.pt)

def show_tables(filename):
    doc = docx.Document(filename)
    for table in doc.tables:
        print(table)

def iter_block_items(parent):
    """
    Yield each paragraph and table child within *parent*, in document order.
    Each returned value is an instance of either Table or Paragraph.
    """
    if isinstance(parent, Document):
        parent_elm = parent._document_part.body._body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child)
        elif isinstance(child, CT_Tbl):
            yield Table(child)


#notes occur multiple times a chapter and thus they combine
#tables arent included
#subtitles are not properly handled due to formatting inconsistencies
def seperate(filename):
    doc = docx.Document(filename)
    dict = {}
    heading = ["" for x in range(len(doc.paragraphs))]
    head_count = -1
    main_heading = ["" for x in range(len(doc.paragraphs))]
    main_h_count = -1

    #for each paragraph
    for i in range(len(doc.paragraphs)):

        #check if current paragraph is empty
        if len(doc.paragraphs[i].runs) == 0:
            continue

        #its a heading
        elif (doc.paragraphs[i].runs[0].bold):
            if (doc.paragraphs[i].runs[0].text.startswith('V') and ("-" in doc.paragraphs[i].runs[0].text )):
                #it is a big heading
                main_h_count = main_h_count + 1
                for j in range(len(doc.paragraphs[i].runs)):
                    if(doc.paragraphs[i].runs[j].bold):
                        main_heading[main_h_count] = main_heading[main_h_count] + doc.paragraphs[i].runs[j].text
                main_heading[main_h_count] = main_heading[main_h_count] + ":"



            else:
                #its a smol heading
                head_count = head_count + 1
                for j in range(len(doc.paragraphs[i].runs)):
                    if(doc.paragraphs[i].runs[j].bold):
                        heading[head_count] = heading[head_count] + doc.paragraphs[i].runs[j].text
                    else: #if the rest of it isnt bold its text!
                        key = main_heading[main_h_count] + heading[head_count]
                        if key in dict:
                            dict[key] = dict.get(key) + doc.paragraphs[i].text
                        else:
                            dict[key] = doc.paragraphs[i].text



        else:
            #no its just text that isnt empty add to its  with a main heading for uniqueness
            key =  main_heading[main_h_count] + heading[head_count]
            if key in dict:
                dict[key] = dict.get(key) + doc.paragraphs[i].text
            else:
                dict[key] = doc.paragraphs[i].text

    #prints the keys and their corresponding text
    for a in dict:
        print("KEY")
        print(a)
        print('\n')
        print('\n')
        print("TEXT")
        print(dict[a])
        print('\n')
        print('\n')



#file_name = 'SCM_sample.docx'
#show_breaks(file_name)
#iter_block_items(file_name)
#print_paragraphs(file_name)
#show_runs(file_name)
#show_font_size(file_name)
#show_tables(file_name)
file_name = 'SCM.docx'
seperate(file_name)
